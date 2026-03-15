from flask import Flask, request, jsonify
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
import PyPDF2
import docx
import nltk
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from google import genai
import os
import re
import json
from datetime import datetime
import secrets
from werkzeug.security import generate_password_hash, check_password_hash
import time

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://resumeapp:123456@localhost:3306/resume_analyzer'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

try:
    client = genai.Client(api_key=os.getenv('GEMINI_API_KEY', 'YOUR_API_KEY'))
    AI_AVAILABLE = True
except:
    client = None
    AI_AVAILABLE = False


class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    email = db.Column(db.String(255), unique=True, nullable=False, index=True)
    name = db.Column(db.String(255), nullable=False)
    password = db.Column(db.String(500), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    analyses = db.relationship('Analysis', backref='user', lazy=True, cascade='all, delete-orphan')
    sessions = db.relationship('Session', backref='user', lazy=True, cascade='all, delete-orphan')

class Session(db.Model):
    __tablename__ = 'sessions'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    token = db.Column(db.String(255), unique=True, nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Analysis(db.Model):
    __tablename__ = 'analyses'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    analysis_id = db.Column(db.String(64), unique=True, nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    company_name = db.Column(db.String(255), nullable=True)
    overall_score = db.Column(db.Integer, nullable=True)
    job_match_score = db.Column(db.Float, nullable=True)
    experience_level = db.Column(db.String(50), nullable=True)
    result_json = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class DatabaseManager:
    @staticmethod
    def create_user(email, name, password):
        try:
            u = User(email=email.lower().strip(), name=name.strip(), password=generate_password_hash(password))
            db.session.add(u); db.session.commit(); return u
        except Exception as e: db.session.rollback(); raise e

    @staticmethod
    def get_user_by_email(email):
        return User.query.filter_by(email=email.lower().strip()).first()

    @staticmethod
    def create_session(user_id):
        try:
            t = secrets.token_hex(32); s = Session(token=t, user_id=user_id)
            db.session.add(s); db.session.commit(); return t
        except Exception as e: db.session.rollback(); raise e

    @staticmethod
    def get_user_by_token(token):
        s = Session.query.filter_by(token=token).first()
        return s.user if s else None

    @staticmethod
    def delete_session(token):
        try:
            s = Session.query.filter_by(token=token).first()
            if s: db.session.delete(s); db.session.commit()
        except: db.session.rollback()

    @staticmethod
    def create_analysis(user_id, company_name, result):
        try:
            a = Analysis(analysis_id=secrets.token_hex(16), user_id=user_id, company_name=company_name,
                overall_score=result.get('overallScore', 0), job_match_score=result.get('jobMatchScore', 0),
                experience_level=result.get('experienceLevel', 'Unknown'), result_json=json.dumps(result))
            db.session.add(a); db.session.commit(); return a
        except Exception as e: db.session.rollback(); raise e


# ============================================================================
# SOFT SKILLS, KNOWN SKILLS, VARIATIONS, STRICT, CONTEXT, FALSE POSITIVES
# (same as previous version - no changes)
# ============================================================================
SOFT_SKILLS = {
    'communication skills', 'communication', 'verbal communication', 'written communication',
    'presentation skills', 'presentations', 'presenting', 'public speaking',
    'interpersonal skills', 'interpersonal', 'stakeholder engagement', 'stakeholder management',
    'client management', 'customer focus', 'cross-functional collaboration', 'collaboration',
    'teamwork', 'team player', 'leadership', 'team leadership', 'tech lead',
    'team management', 'people management', 'mentoring', 'mentorship', 'coaching',
    'decision making', 'decision-making', 'strategic thinking', 'problem solving',
    'problem-solving', 'troubleshooting', 'critical thinking', 'analytical thinking',
    'creative thinking', 'innovation mindset', 'innovative thinking', 'innovation',
    'learning agility', 'fast learner', 'quick learner', 'adaptability',
    'resilience', 'resilient', 'ownership', 'sense of ownership', 'takes ownership',
    'attention to detail', 'detail oriented', 'detail-oriented', 'time management',
    'self-motivated', 'proactive', 'initiative', 'coordination', 'organization',
    'multitasking', 'prioritization', 'flexibility', 'work ethic', 'professionalism',
    'reliability', 'accountability', 'integrity', 'conflict resolution', 'negotiation',
    'persuasion', 'emotional intelligence', 'empathy', 'active listening',
    'business acumen', 'strategic planning', 'goal setting', 'vendor management',
    'procurement', 'budget management', 'quality focus',
}

KNOWN_SKILLS = [
    'python','java','javascript','typescript','c++','c#','ruby','php','swift','kotlin','go','rust',
    'scala','perl','r','matlab','dart','lua','haskell','elixir','clojure','erlang','f#','objective-c',
    'cobol','fortran','assembly','groovy','julia','solidity','visual basic','vba','shell scripting','apex',
    'react','angular','vue','svelte','next.js','nuxt','gatsby','htmx','remix','astro',
    'node.js','express','nest.js','django','flask','fastapi','spring boot','laravel','rails','asp.net','.net',
    'gin','fiber','actix','rocket','phoenix','koa','hapi','strapi',
    'sql','mysql','postgresql','mongodb','redis','oracle','sqlite','dynamodb','cassandra','elasticsearch',
    'firebase','supabase','mariadb','couchdb','neo4j','influxdb','timescaledb','cockroachdb','snowflake',
    'bigquery','redshift','memcached','planetscale',
    'aws','azure','gcp','heroku','vercel','netlify','digitalocean','cloudflare',
    'docker','kubernetes','jenkins','gitlab ci','github actions','terraform','ansible','prometheus','grafana',
    'nginx','apache','traefik','puppet','vagrant','argocd','istio','envoy',
    'datadog','new relic','splunk','elk stack','logstash','kibana','fluentd','jaeger','zipkin','opentelemetry',
    'circleci','travis ci','bamboo','teamcity','pulumi','rancher',
    'machine learning','deep learning','tensorflow','pytorch','keras','scikit-learn','pandas','numpy','opencv',
    'nlp','langchain','hugging face','transformers','bert','gpt','llm','computer vision','neural networks',
    'reinforcement learning','generative ai','prompt engineering','rag','mlops','mlflow','kubeflow','sagemaker',
    'vertex ai','xgboost','lightgbm','catboost','natural language processing','sentiment analysis',
    'object detection','image classification','chatbot development','conversational ai',
    'data analysis','data visualization','data modeling','data engineering','data warehousing','data mining',
    'data pipeline','data lake','etl','elt','business intelligence','predictive analytics','statistical analysis',
    'power bi','tableau','looker','metabase','superset','qlik','google data studio',
    'microsoft excel','google sheets','jupyter notebook','google colab','databricks',
    'apache spark','apache flink','apache beam','apache airflow','prefect','dagster','dbt',
    'android','ios','react native','flutter','xamarin','ionic','swiftui','jetpack compose','kotlin multiplatform',
    'html','css','sass','tailwind','bootstrap','material-ui','webpack','vite','babel','rollup','esbuild','parcel',
    'websocket','webrtc','web assembly','seo','web accessibility','styled-components','chakra ui','ant design',
    'graphql','rest api','grpc','oauth','jwt','swagger','postman','api gateway','webhook','soap','protobuf',
    'jest','cypress','selenium','pytest','unit testing','integration testing','e2e testing','tdd','bdd',
    'mocha','chai','jasmine','playwright','puppeteer','junit','testng','mockito',
    'jmeter','k6','gatling','locust','test automation','robot framework','appium','storybook',
    'git','github','gitlab','bitbucket','jira','confluence','trello','asana','figma','sketch','adobe xd',
    'microservices','serverless','event-driven architecture','domain-driven design','clean architecture',
    'mvc','mvvm','design patterns','solid principles','system design','distributed systems',
    'rabbitmq','kafka','celery','apache pulsar','amazon sqs','amazon sns',
    'cybersecurity','owasp','ssl/tls','sso','saml','ldap','active directory','mfa','iam','siem',
    'zero trust','devsecops','penetration testing','soc2','gdpr','hipaa','pci dss',
    'linux','bash','powershell','ubuntu','centos','rhel','debian','windows server',
    'blockchain','ethereum','smart contracts','web3','hardhat','truffle',
    'unity','unreal engine','godot','opengl','directx','vulkan','blender','maya',
    'sap','salesforce','oracle erp','microsoft dynamics','workday','servicenow',
    'agile','scrum','kanban','prince2','pmp',
    'business analysis','requirements gathering','brd','frd','srs','use cases','user stories',
    'gap analysis','process mapping','uat','user acceptance testing','root cause analysis',
    'microsoft office','microsoft word','microsoft powerpoint','microsoft visio','microsoft project','google workspace',
    'ci/cd','continuous integration','continuous deployment','infrastructure as code','gitops',
    'iot','arduino','raspberry pi','esp32','mqtt','edge computing','embedded systems',
    'power apps','power automate','zapier','retool','bubble','webflow','outsystems','mendix',
    'rpa','uipath','automation anywhere','blue prism',
    'wordpress','drupal','contentful','sentry','bugsnag','pagerduty',
    'performance optimization','query optimization','lazy loading','code splitting',
    'responsive design','progressive web apps','technical writing',
]

SKILL_VARIATIONS = {
    'react':['react','reactjs','react.js'],'vue':['vue','vuejs','vue.js'],
    'angular':['angular','angularjs','angular.js'],'node.js':['node.js','nodejs'],
    'next.js':['next.js','nextjs'],'nest.js':['nest.js','nestjs'],
    'express':['express.js','expressjs'],'typescript':['typescript'],
    'javascript':['javascript','es6','ecmascript'],'python':['python','python3'],
    'c++':['c++','cpp','cplusplus'],'c#':['c#','csharp','c sharp'],
    '.net':['.net','dotnet','asp.net'],'postgresql':['postgresql','postgres','psql'],
    'mongodb':['mongodb','mongo'],'kubernetes':['kubernetes','k8s'],
    'aws':['aws','amazon web services'],'gcp':['gcp','google cloud','google cloud platform'],
    'azure':['azure','microsoft azure'],'machine learning':['machine learning'],
    'deep learning':['deep learning'],'tensorflow':['tensorflow'],
    'scikit-learn':['scikit-learn','sklearn','scikit learn'],
    'react native':['react native','react-native','reactnative'],
    'spring boot':['spring boot','springboot','spring-boot','spring framework','spring mvc','spring cloud','spring security'],
    'rest api':['rest api','restful api','rest apis','restful apis','restful services','restful'],
    'ci/cd':['ci/cd','cicd','ci cd','ci-cd'],'github actions':['github actions'],
    'gitlab ci':['gitlab ci','gitlab-ci','gitlab ci/cd'],
    'power bi':['power bi','powerbi','power-bi','microsoft power bi'],
    'tableau':['tableau','tableau desktop','tableau server'],
    'microsoft excel':['microsoft excel','ms excel','advanced excel','excel vba','excel macro','excel pivot'],
    'google sheets':['google sheets'],
    'data analysis':['data analysis','data analytics'],'data visualization':['data visualization','data viz'],
    'business intelligence':['business intelligence'],
    'predictive analytics':['predictive analytics','predictive modeling'],
    'etl':['etl','extract transform load'],'data pipeline':['data pipeline','data pipelines'],
    'data warehousing':['data warehousing','data warehouse'],'data lake':['data lake','datalake'],
    'apache spark':['apache spark','pyspark'],'apache airflow':['apache airflow','airflow'],
    'dbt':['dbt','data build tool'],'snowflake':['snowflake'],
    'bigquery':['bigquery','big query','google bigquery'],'redshift':['redshift','amazon redshift'],
    'databricks':['databricks'],'business analysis':['business analysis'],
    'requirements gathering':['requirements gathering','requirement gathering'],
    'brd':['brd','business requirements document'],'frd':['frd','functional requirements document'],
    'user stories':['user stories','user story'],'use cases':['use cases','use case'],
    'uat':['uat','user acceptance testing'],'gap analysis':['gap analysis','fit-gap analysis'],
    'process mapping':['process mapping','process flow'],'root cause analysis':['root cause analysis','rca'],
    'microsoft office':['microsoft office','ms office','office 365'],
    'microsoft word':['microsoft word','ms word'],
    'microsoft powerpoint':['microsoft powerpoint','ms powerpoint','powerpoint'],
    'microsoft visio':['microsoft visio','ms visio'],
    'agile':['agile methodology','agile development','agile framework','agile practices','agile scrum'],
    'scrum':['scrum','scrum master','scrum framework'],'kanban':['kanban'],
    'cybersecurity':['cybersecurity','cyber security','information security','infosec'],
    'owasp':['owasp','owasp top 10'],'ssl/tls':['ssl/tls','ssl','tls'],
    'sso':['sso','single sign-on'],'mfa':['mfa','multi-factor authentication','2fa'],
    'iam':['iam','identity and access management'],'gdpr':['gdpr'],'soc2':['soc2','soc 2'],
    'devsecops':['devsecops'],'elk stack':['elk stack','elk','elastic stack'],
    'infrastructure as code':['infrastructure as code','iac'],'gitops':['gitops'],
    'argocd':['argocd','argo cd','argo-cd'],
    'unit testing':['unit testing','unit tests'],'integration testing':['integration testing','integration tests'],
    'e2e testing':['e2e testing','end-to-end testing','end to end testing'],
    'tdd':['tdd','test driven development'],'bdd':['bdd','behavior driven development'],
    'test automation':['test automation','automated testing','automation testing'],
    'penetration testing':['penetration testing','pen testing','pentest','ethical hacking'],
    'microservices':['microservices','microservice','micro-services'],
    'serverless':['serverless','faas'],
    'event-driven architecture':['event-driven architecture','event driven architecture'],
    'domain-driven design':['domain-driven design','domain driven design'],
    'system design':['system design','systems design'],
    'distributed systems':['distributed systems','distributed computing'],
    'design patterns':['design patterns','design pattern'],'solid principles':['solid principles'],
    'clean architecture':['clean architecture'],
    'rpa':['rpa','robotic process automation'],'uipath':['uipath'],
    'automation anywhere':['automation anywhere'],'blue prism':['blue prism'],
    'power automate':['power automate','microsoft power automate'],
    'iot':['iot','internet of things'],'raspberry pi':['raspberry pi'],
    'edge computing':['edge computing'],'embedded systems':['embedded systems'],
    'power apps':['power apps','powerapps'],
    'sap':['sap','sap erp','sap s/4hana'],'salesforce':['salesforce','sfdc'],
    'servicenow':['servicenow','service now'],
    'blockchain':['blockchain'],'ethereum':['ethereum'],
    'smart contracts':['smart contracts','smart contract'],'web3':['web3','web 3.0'],
    'unity':['unity engine','unity 3d','unity game','unity developer'],
    'blender':['blender 3d','blender modeling','blender animation'],
    'maya':['autodesk maya','maya 3d','maya modeling'],
    'sentry':['sentry','sentry.io'],'datadog':['datadog'],'new relic':['new relic','newrelic'],
    'performance optimization':['performance optimization','performance tuning'],
    'query optimization':['query optimization','sql optimization','query tuning'],
    'lazy loading':['lazy loading'],'code splitting':['code splitting'],
    'generative ai':['generative ai','genai','gen ai'],'prompt engineering':['prompt engineering'],
    'llm':['llm','large language model','large language models'],
    'rag':['rag','retrieval augmented generation'],
    'hugging face':['hugging face','huggingface'],'chatbot development':['chatbot development','chatbot'],
    'mlops':['mlops','ml ops'],
    'responsive design':['responsive design','responsive web design'],
    'progressive web apps':['progressive web app','progressive web apps','pwa'],
    'wordpress':['wordpress'],'web accessibility':['web accessibility','a11y'],
    'seo':['seo','search engine optimization'],'swagger':['swagger','openapi'],'postman':['postman'],
    'technical writing':['technical writing','tech writing'],
}

STRICT_MATCH_SKILLS = {
    'go','r','c','dart','ruby','rust','swift','scala','perl','express','flask','rails',
    'ionic','redux','helm','kafka','nest.js','nuxt','svelte','grpc','oauth','jwt','vite',
    'babel','jest','cypress','celery','bash','gin','fiber','rocket','phoenix','koa','hapi',
    'sass','consul','vault','flux','nats','expo','remix','astro','apex','lua','julia',
    'agile','node.js','sentry','mqtt','sql','css','html','unity','blender','maya','puppet','vagrant',
}

CONTEXT_REQUIRED_SKILLS = {
    'go','scala','rust','swift','dart','ruby','perl','express','flask','rails',
    'phoenix','rocket','gin','fiber','koa','hapi','consul','vault','flux','nats','expo','remix',
    'agile','apex','lua','julia','helm','celery','babel','sass','ionic','redux','kafka','mqtt',
    'unity','blender','maya','puppet','vagrant',
}

FALSE_POSITIVE_PATTERNS = {
    'go':['go-live','golive','go live','going','good','goal','goals','gone','got','govern','governance','google','goods','ago','undergo','ego','ergo','logo','mango','category','cargo','ongoing','outgoing','forgo'],
    'r':['brd','frd','crp','uat','erp','mr','dr','sr','jr','pr','or','for','are','our','your','their','her','per','after','under','over','other','every','where','there','here','more','before','never','ever','during','from'],
    'scala':['scalable','scalability','scaling','escalate','escalating','escalation','scale','scaled'],
    'rust':['trust','trusted','trustworthy','robust','industry','frustrate','frustrated','entrust'],
    'swift':['swiftly','swiftness'],'dart':['standard','standards','darting','dashboard'],
    'ruby':['rubric','rubrics'],
    'express':['express interest','express ideas','expressed','expressing','expression','expressly','express their','express your','express the','express concern','express delivery'],
    'flask':['flashback'],'rails':['trail','trails','derail','derails','guardrails'],
    'helm':['help','helped','helping','helpful','overwhelm','overwhelming'],
    'ionic':['electronic','electronics','chronicle','chronically'],
    'sass':['saas','sas'],'redux':['reduce','reduced','reducing','reduction'],
    'perl':['experience','expert','expertise','preferably','properly','superlative','interpersonal'],
    'jest':['suggest','suggestion','majestic','majesty'],'celery':['salary','accelery'],
    'vite':['invite','invited','invitation','favorite','favourite'],
    'bash':['bashing','abash'],'consul':['consult','consultant','consulting','consultation'],
    'vault':['default','defaults'],'flux':['influx'],'nats':['natural','naturally','national'],
    'expo':['expose','exposed','exposure','export','exported'],
    'gin':['going','beginning','engineering','managing','changing','emerging','designing','login','margin','origin'],
    'fiber':['fibers'],'hapi':['happy','happiness','happily'],
    'lua':['value','evaluate','evaluation','valuable'],
    'css':['access','accessing','success','successful','necessary'],
    'puppet':['puppet show','puppeteer'],'unity':['community','opportunity','unit'],
    'blender':['blended','blending'],'vagrant':['vagrancy'],
    'agile':[],'phoenix':[],'koa':[],'rocket':[],'remix':[],'astro':[],
    'apex':[],'julia':[],'kafka':[],'nuxt':[],'svelte':[],
    'grpc':[],'oauth':[],'jwt':[],'babel':[],'cypress':[],
    'nest.js':[],'node.js':[],'sentry':[],'mqtt':[],'sql':[],'html':[],'maya':[],
}

TECH_CONTEXT_WORDS = [
    'programming','language','framework','library','developer','development','software','code','coding',
    'engineer','engineering','stack','backend','frontend','fullstack','full-stack','api','database','server',
    'cloud','deploy','deployment','devops','container','microservice','application','script','scripting',
    'testing','debug','compile','compiler','runtime','sdk','ide','package','algorithm','machine learning',
    'ai','ml','repository','version control','ci/cd','build','release','infrastructure','kubernetes',
    'docker','aws','azure','gcp','proficient','expertise','skilled','technologies','tools','platform',
    'architecture','experience with','knowledge of','proficiency','tech stack','technical','integration',
    'react','angular','vue','python','java','node','django','fastapi','spring boot',
    'terraform','jenkins','git','github','gitlab',
]

RELATED_SKILL_GROUPS = {
    'frontend':['react','angular','vue','svelte','next.js','nuxt','gatsby','remix','astro'],
    'backend':['express','django','flask','fastapi','spring boot','nest.js','laravel','rails','gin','fiber','phoenix','koa'],
    'languages':['python','java','javascript','typescript','c++','c#','go','rust','ruby','kotlin','swift','php'],
    'sql_db':['sql','mysql','postgresql','oracle','sqlite','mariadb','snowflake','bigquery','redshift'],
    'nosql_db':['mongodb','redis','dynamodb','cassandra','firebase','couchdb','neo4j'],
    'cloud':['aws','azure','gcp'],
    'containers':['docker','kubernetes','rancher'],
    'ci_cd':['jenkins','gitlab ci','github actions','ci/cd','circleci','travis ci','bamboo','teamcity'],
    'iac':['terraform','ansible','puppet','pulumi'],
    'monitoring':['prometheus','grafana','datadog','new relic','splunk','elk stack','jaeger','zipkin'],
    'css_fw':['css','sass','tailwind','bootstrap','material-ui','chakra ui','ant design','styled-components'],
    'testing':['jest','cypress','selenium','pytest','playwright','puppeteer','junit','testng','mocha'],
    'mobile':['android','ios','react native','flutter','xamarin','ionic','swiftui','jetpack compose'],
    'ml':['machine learning','deep learning','tensorflow','pytorch','keras','scikit-learn','xgboost','lightgbm'],
    'data':['pandas','numpy','r','matlab','jupyter notebook','google colab'],
    'data_proc':['apache spark','apache flink','apache beam','apache airflow','dbt','prefect','dagster'],
    'bi':['power bi','tableau','looker','metabase','superset','qlik','google data studio'],
    'messaging':['rabbitmq','kafka','celery','amazon sqs','amazon sns','apache pulsar'],
    'api':['rest api','graphql','grpc','websocket','soap'],
    'vcs':['git','github','gitlab','bitbucket'],
    'auth':['oauth','jwt','sso','saml','ldap'],
    'pm':['jira','confluence','trello','asana'],
    'design':['figma','sketch','adobe xd'],
    'security':['cybersecurity','owasp','penetration testing','devsecops','siem'],
    'methodology':['agile','scrum','kanban'],
    'ba':['business analysis','requirements gathering','gap analysis','process mapping','uat','brd','frd','user stories','use cases'],
    'erp':['sap','salesforce','oracle erp','microsoft dynamics','workday','servicenow'],
    'office':['microsoft office','microsoft excel','microsoft word','microsoft powerpoint','google workspace','google sheets'],
    'rpa':['rpa','uipath','automation anywhere','blue prism','power automate'],
    'blockchain':['blockchain','ethereum','solidity','smart contracts','web3'],
    'ai_adv':['generative ai','prompt engineering','llm','rag','langchain','hugging face','chatbot development'],
    'arch':['microservices','serverless','event-driven architecture','domain-driven design','system design','distributed systems','clean architecture','design patterns'],
    'spreadsheets':['microsoft excel','google sheets'],
    'iot':['iot','arduino','raspberry pi','mqtt','edge computing','embedded systems'],
    'lowcode':['power apps','retool','bubble','webflow','outsystems','mendix'],
}

KNOWN_CERTIFICATIONS = [
    'aws certified','aws solutions architect','aws developer','aws sysops','aws cloud practitioner',
    'aws devops engineer','aws machine learning','azure certified','azure fundamentals','az-900','az-104',
    'az-204','az-400','google cloud certified','gcp professional','google cloud engineer',
    'certified kubernetes administrator','cka','ckad','cks','terraform certified','hashicorp certified',
    'pmp','project management professional','prince2','prince2 certified',
    'scrum master','csm','psm','cissp','cism','ceh','certified ethical hacker',
    'comptia security+','comptia network+','comptia a+','ccna','ccnp','cisco certified',
    'oracle certified','ocp','oca','microsoft certified','mcsa','mcse',
    'salesforce certified','salesforce administrator','sap certified',
    'istqb','istqb certified','itil certified','itil foundation',
    'rhce','red hat certified','docker certified','dca',
    'certified scrum developer','csd','togaf certified',
    'six sigma','six sigma green belt','six sigma black belt',
    'tensorflow developer certificate','tensorflow certified',
    'databricks certified','snowflake certified','mongodb certified',
    'elastic certified','linux foundation certified','lfcs','lfce',
    'certified information systems auditor','cisa','google analytics certified',
    'meta certified',
]

def safe_str(v, d=""): return str(v) if v is not None else d
def safe_list(v): return v if isinstance(v, list) else []
def safe_dict(v): return v if isinstance(v, dict) else {}
def safe_int(v, d=0):
    try: return int(v)
    except: return d
def is_soft_skill(s): return s.lower().strip() in SOFT_SKILLS if s else False


def get_skill_resources_data(skill):
    if not skill: skill = "Unknown"
    dn = skill.title()
    data = {
        'displayName': dn, 'category': 'Technical Skill', 'difficulty': 'Medium',
        'learningTime': '4-6 weeks', 'salaryImpact': '+$10K-$20K', 'demandLevel': 'High',
        'description': f'{dn} is a valuable technical skill.',
        'careerPaths': ['Software Developer'], 'relatedSkills': [],
        'freeCoursePlatforms': [
            {'platform': 'freeCodeCamp', 'courseName': f'Learn {dn}', 'duration': '10-20h', 'level': 'Beginner', 'link': f'https://www.youtube.com/results?search_query=freecodecamp+{skill}', 'rating': '4.8/5'},
            {'platform': 'Udemy', 'courseName': f'{dn} Guide', 'duration': '20-40h', 'level': 'All', 'link': f'https://www.udemy.com/courses/search/?q={skill}', 'rating': '4.6/5'},
        ],
        'youtubeChannels': [
            {'channelName': 'freeCodeCamp.org', 'videoTitle': f'{dn} Course', 'duration': '4-8h', 'subscribers': '9M+', 'link': f'https://www.youtube.com/results?search_query=freecodecamp+{skill}'},
        ],
        'practiceWebsites': [{'platform': 'LeetCode', 'description': f'{dn} practice', 'difficulty': 'Medium-Hard', 'link': 'https://leetcode.com', 'type': 'Practice'}],
        'projectIdeas': [
            {'projectName': f'{dn} Starter', 'difficulty': 'Beginner', 'duration': '1 week', 'description': f'Learn {dn}', 'skills': [dn], 'whatYouLearn': f'Core {dn}'},
            {'projectName': f'{dn} Portfolio', 'difficulty': 'Intermediate', 'duration': '2 weeks', 'description': f'Showcase {dn}', 'skills': [dn], 'whatYouLearn': f'{dn} patterns'},
        ],
        'learningRoadmap': {'week1': f'{dn} fundamentals', 'week2': 'First project', 'week3': 'Intermediate', 'week4': 'Portfolio'},
        'certifications': []
    }
    return data


# ============================================================================
# RESUME ANALYZER WITH 10-METRIC SCORING + DETAILED BREAKDOWNS
# ============================================================================
class ResumeAnalyzer:
    def __init__(self, ai_client=None):
        self.client = ai_client
        self.vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
        self.month_map = {'jan':1,'january':1,'feb':2,'february':2,'mar':3,'march':3,'apr':4,'april':4,
            'may':5,'jun':6,'june':6,'jul':7,'july':7,'aug':8,'august':8,'sep':9,'sept':9,'september':9,
            'oct':10,'october':10,'nov':11,'november':11,'dec':12,'december':12}

    def extract_pdf_text(self, file):
        try:
            reader = PyPDF2.PdfReader(file)
            return "\n".join([(p.extract_text() or "") for p in reader.pages])
        except: return ""

    def extract_docx_text(self, file):
        try:
            doc = docx.Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells: text += " " + c.text
            return text
        except: return ""

    def is_false_positive(self, skill, text_lower, pos):
        for fp in FALSE_POSITIVE_PATTERNS.get(skill.lower(), []):
            if fp in text_lower:
                for m in re.finditer(re.escape(fp), text_lower):
                    if m.start() <= pos < m.start() + len(fp): return True
        return False

    def skill_exists_strict(self, skill, text):
        text_lower = text.lower(); skill_lower = skill.lower()
        variations = SKILL_VARIATIONS.get(skill_lower, [skill_lower])
        is_strict = skill_lower in STRICT_MATCH_SKILLS
        needs_ctx = skill_lower in CONTEXT_REQUIRED_SKILLS
        for var in variations:
            if len(var) <= 2:
                if self._match_single_letter(var, skill_lower, text_lower): return True
                continue
            pattern = r'\b' + re.escape(var) + r'\b'
            matches = list(re.finditer(pattern, text_lower))
            if not matches: continue
            for match in matches:
                pos, end_pos = match.start(), match.end()
                if self.is_false_positive(skill_lower, text_lower, pos): continue
                if is_strict:
                    cb = text_lower[pos-1] if pos > 0 else ' '
                    ca = text_lower[end_pos] if end_pos < len(text_lower) else ' '
                    if cb.isalpha() or ca.isalpha(): continue
                if needs_ctx:
                    ctx = text_lower[max(0,pos-80):min(len(text_lower),end_pos+80)]
                    if any(tw in ctx for tw in TECH_CONTEXT_WORDS) or any(p in ctx for p in [',','/','|','•','·','skills:','technologies:','tech stack:','proficient in','experience with','knowledge of','familiar with','worked with','hands-on','expertise in','required:','requirements:','qualifications:','must have','nice to have']): return True
                    continue
                return True
        return False

    def _match_single_letter(self, var, skill_name, text_lower):
        pats = {'r':[r'\br\s+programming',r'\br\s+language',r'\br\s+studio',r'rstudio',r'programming\s+in\s+r\b',r'language:\s*r\b',r'r,\s*(?:python|matlab|sas|spss)',r'(?:python|matlab|sas|spss),\s*r\b',r'\br\s+(?:shiny|tidyverse|ggplot|dplyr|cran)'],
                'c':[r'\bc\s+programming',r'\bc\s+language',r'programming\s+in\s+c\b',r'\bc/c\+\+',r'\bc,\s*c\+\+',r'language:\s*c\b']}
        for p in pats.get(skill_name, []):
            if re.search(p, text_lower): return True
        return False

    def extract_skills_from_text(self, text):
        found = []
        for skill in KNOWN_SKILLS:
            if is_soft_skill(skill): continue
            if self.skill_exists_strict(skill, text):
                if skill not in found: found.append(skill)
        return found

    def match_skills(self, job_skills, resume_text):
        matched, missing, skipped = [], [], []
        for skill in job_skills:
            if is_soft_skill(skill): skipped.append(skill); continue
            if self.skill_exists_strict(skill, resume_text): matched.append(skill)
            else: missing.append(skill)
        return matched, missing

    # ===================================================================
    # DETAILED EXTRACTION METHODS (for tap-to-view)
    # ===================================================================

    def get_skill_match_details(self, job_skills, resume_text):
        """Returns detailed info for each skill - matched/missing + where found"""
        details = []
        tl = resume_text.lower()
        for skill in job_skills:
            if is_soft_skill(skill): continue
            found = self.skill_exists_strict(skill, resume_text)
            context_snippet = ""
            if found:
                variations = SKILL_VARIATIONS.get(skill.lower(), [skill.lower()])
                for var in variations:
                    if len(var) <= 2: continue
                    m = re.search(r'\b' + re.escape(var) + r'\b', tl)
                    if m:
                        start = max(0, m.start() - 60)
                        end = min(len(tl), m.end() + 60)
                        context_snippet = "..." + resume_text[start:end].strip().replace('\n', ' ') + "..."
                        break
            details.append({
                "skill": skill,
                "found": found,
                "status": "✅ Matched" if found else "❌ Missing",
                "contextInResume": context_snippet if found else "Not found in resume"
            })
        return details

    def get_skill_depth_details(self, matched_skills, resume_text):
        """Returns how deeply each matched skill is used"""
        tl = resume_text.lower()
        project_words = ['built','developed','created','designed','implemented','architected','deployed',
                        'migrated','integrated','optimized','led','managed','delivered','launched',
                        'automated','project','application','system','platform','service']
        details = []
        for skill in matched_skills:
            variations = SKILL_VARIATIONS.get(skill.lower(), [skill.lower()])
            total_count = 0
            in_project = False
            snippets = []
            for var in variations:
                if len(var) <= 2: continue
                for m in re.finditer(r'\b' + re.escape(var) + r'\b', tl):
                    total_count += 1
                    ctx_s = max(0, m.start() - 80)
                    ctx_e = min(len(tl), m.end() + 80)
                    ctx = tl[ctx_s:ctx_e]
                    if any(pw in ctx for pw in project_words):
                        in_project = True
                    if len(snippets) < 3:
                        snippets.append("..." + resume_text[ctx_s:ctx_e].strip().replace('\n',' ') + "...")

            if in_project and total_count >= 3: depth = "Expert"
            elif total_count >= 3: depth = "Experienced"
            elif in_project: depth = "Practical"
            elif total_count >= 2: depth = "Familiar"
            else: depth = "Mentioned"

            score = 3 if depth in ["Expert"] else 2.5 if depth == "Experienced" else 2 if depth == "Practical" else 1.5 if depth == "Familiar" else 1
            details.append({
                "skill": skill,
                "mentionCount": total_count,
                "usedInProject": in_project,
                "depthLevel": depth,
                "depthScore": f"{score}/3",
                "contextSnippets": snippets
            })
        avg = sum(d['mentionCount'] for d in details) / max(len(details), 1)
        return details, round(avg, 1)

    def get_transferable_details(self, missing_skills, resume_skills):
        """Returns which resume skills are related to each missing skill"""
        details = []
        for ms in missing_skills:
            ml = ms.lower()
            related = []
            group_name = ""
            for gn, gs in RELATED_SKILL_GROUPS.items():
                if ml in gs:
                    group_name = gn
                    for rs in resume_skills:
                        if rs.lower() in gs and rs.lower() != ml:
                            related.append(rs)
            details.append({
                "missingSkill": ms,
                "skillGroup": group_name,
                "relatedSkillsYouHave": related,
                "transferabilityLevel": "High" if len(related) >= 3 else "Medium" if len(related) >= 1 else "None",
                "recommendation": f"Your {', '.join(related[:3])} experience transfers to {ms}" if related else f"No related skills found - start learning {ms} from scratch"
            })
        return details

    def get_tfidf_details(self, resume_text, job_desc):
        """Returns TF-IDF cosine similarity with top matching terms"""
        try:
            vectorizer = TfidfVectorizer(stop_words='english', max_features=200, ngram_range=(1, 2))
            tfidf_matrix = vectorizer.fit_transform([resume_text, job_desc])
            sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            feature_names = vectorizer.get_feature_names_out()
            resume_vec = tfidf_matrix[0].toarray()[0]
            job_vec = tfidf_matrix[1].toarray()[0]
            shared_terms = []
            for i, name in enumerate(feature_names):
                if resume_vec[i] > 0 and job_vec[i] > 0:
                    combined = resume_vec[i] + job_vec[i]
                    shared_terms.append({"term": name, "relevance": round(combined, 3)})
            shared_terms.sort(key=lambda x: x['relevance'], reverse=True)
            top_resume = []
            for i, name in enumerate(feature_names):
                if resume_vec[i] > 0.1:
                    top_resume.append({"term": name, "weight": round(resume_vec[i], 3)})
            top_resume.sort(key=lambda x: x['weight'], reverse=True)
            top_job = []
            for i, name in enumerate(feature_names):
                if job_vec[i] > 0.1:
                    top_job.append({"term": name, "weight": round(job_vec[i], 3)})
            top_job.sort(key=lambda x: x['weight'], reverse=True)
            return {
                "similarityPercent": round(sim * 100, 1),
                "topSharedTerms": shared_terms[:15],
                "topResumeTerms": top_resume[:10],
                "topJobTerms": top_job[:10],
                "interpretation": "Strong alignment" if sim > 0.3 else "Moderate alignment" if sim > 0.15 else "Weak alignment - tailor resume"
            }
        except:
            return {"similarityPercent": 0, "topSharedTerms": [], "topResumeTerms": [], "topJobTerms": [], "interpretation": "Could not compute"}

    def get_keyword_details(self, resume_text, job_desc):
        """Returns exactly which keywords matched and which didn't"""
        try: stop = set(stopwords.words('english'))
        except: stop = set()
        stop.update(['experience','work','team','company','position','role','required','preferred','ability','strong','knowledge','skills','years','working','including','using','etc','responsibilities','requirements','qualifications','job','must','will','also','well','good','great','new','looking','seeking','ideal','candidate','apply','go','live','going','gone','rest','express','spring','node','make','build','excel','safe','lean','leadership','communication','teamwork','collaboration','motivated','proactive','adaptable','flexible','creative','innovative','detail','oriented','interpersonal','mentoring','coaching','monitor','monitoring','report','reporting','pipeline','optimization','sprint','roadmap','documentation','compliance','networking','edge','embedded','cache','insights','analytics','routing','switching','alerting','logging','tracing','profiling','indexing','blender','unity','notion','linear','slack','discord','puppet','vagrant','render','ghost','lit','echo'])
        jw = [w for w in re.findall(r'[a-zA-Z]+', job_desc.lower()) if len(w) > 2 and w not in stop]
        unique_jw = list(dict.fromkeys(jw))
        rl = resume_text.lower()
        matched_kw = [w for w in unique_jw if w in rl]
        unmatched_kw = [w for w in unique_jw if w not in rl]
        pct = (len(matched_kw) / len(unique_jw) * 100) if unique_jw else 50
        return {
            "totalKeywords": len(unique_jw),
            "matchedCount": len(matched_kw),
            "unmatchedCount": len(unmatched_kw),
            "matchPercent": round(pct, 1),
            "matchedKeywords": matched_kw[:30],
            "unmatchedKeywords": unmatched_kw[:20],
            "suggestion": "Good keyword coverage" if pct >= 60 else "Add more JD keywords to resume" if pct >= 30 else "Resume needs significant keyword optimization"
        }

    def extract_work_experiences(self, text):
        """Extract structured work experience entries"""
        experiences = []
        tl = text.lower()
        cy = datetime.now().year
        cm = datetime.now().month
        mp = r'(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
        pp = r'(?:present|current|now|ongoing|till\s*date|to\s*date|today)'
        ds = r'\s*(?:[-–—]+|\bto\b)\s*'
        title_patterns = [
            r'(?:software|web|frontend|backend|full[\-\s]?stack|senior|junior|lead|staff|principal|associate)\s+(?:engineer|developer|programmer|architect)',
            r'(?:data|ml|ai|cloud|devops|platform|mobile|ios|android|qa|test|automation)\s+(?:engineer|scientist|analyst|developer)',
            r'(?:project|product|engineering|program|delivery)\s+manager',
            r'(?:technical|solutions?|enterprise|software|system)\s+architect',
            r'(?:business|systems?|data|financial|operations?)\s+analyst',
            r'(?:scrum\s+master|tech(?:nical)?\s+lead|team\s+lead|engineering\s+manager)',
            r'(?:intern(?:ship)?|trainee|apprentice)',
            r'(?:consultant|contractor|freelancer)',
            r'(?:director|vp|vice president|head)\s+(?:of\s+)?(?:engineering|technology|product|data)',
            r'cto|chief\s+technology\s+officer',
        ]
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            if not line_lower or len(line_lower) < 5: continue
            for tp in title_patterns:
                m = re.search(tp, line_lower)
                if m:
                    title = m.group(0).strip().title()
                    company = ""
                    dates = ""
                    duration = ""
                    nearby_text = "\n".join(lines[max(0,i-1):min(len(lines),i+4)])
                    nearby_lower = nearby_text.lower()
                    date_m = re.search(rf'({mp})[\.,]?\s*((?:19|20)\d{{2}}){ds}(?:({mp})[\.,]?\s*((?:19|20)\d{{2}})|({pp}))', nearby_lower)
                    if date_m:
                        sy = int(date_m.group(2))
                        sm_name = date_m.group(1)[:3]
                        if date_m.group(5):
                            dates = f"{sm_name.title()} {sy} - Present"
                            months = (cy - sy) * 12 + (cm - self.month_map.get(sm_name, 1))
                            duration = f"{months // 12}y {months % 12}m"
                        elif date_m.group(3) and date_m.group(4):
                            ey = int(date_m.group(4))
                            em_name = date_m.group(3)[:3]
                            dates = f"{sm_name.title()} {sy} - {em_name.title()} {ey}"
                            em = self.month_map.get(em_name, 12)
                            sm = self.month_map.get(sm_name, 1)
                            months = (ey - sy) * 12 + (em - sm)
                            duration = f"{months // 12}y {months % 12}m"
                    at_match = re.search(r'(?:at|@|-|–|,)\s*([A-Z][A-Za-z\s&.,]+)', line)
                    if at_match:
                        company = at_match.group(1).strip()[:50]
                    desc_lines = []
                    for j in range(i+1, min(i+6, len(lines))):
                        l = lines[j].strip()
                        if l and (l.startswith('•') or l.startswith('-') or l.startswith('*') or l[0].islower()):
                            desc_lines.append(l[:150])
                        elif l and any(re.search(tp2, l.lower()) for tp2 in title_patterns):
                            break
                    already = any(e['title'] == title and e.get('dates','') == dates for e in experiences)
                    if not already:
                        experiences.append({
                            "title": title,
                            "company": company,
                            "dates": dates,
                            "duration": duration,
                            "responsibilities": desc_lines[:4]
                        })
                    break
        return experiences

    def extract_education_details(self, text):
        """Extract structured education entries"""
        education = []
        tl = text.lower()
        degree_patterns = [
            (r'(?:ph\.?d|doctorate|doctoral)\s*(?:in|of)?\s*([A-Za-z\s]+)?', 'PhD'),
            (r"(?:master'?s?|m\.?s\.?|m\.?sc\.?|mba|m\.?b\.?a|m\.?tech|m\.?e\.?)\s*(?:in|of)?\s*([A-Za-z\s]+)?", 'Masters'),
            (r"(?:bachelor'?s?|b\.?s\.?|b\.?sc\.?|b\.?tech|b\.?e\.?|b\.?a\.?)\s*(?:in|of)?\s*([A-Za-z\s]+)?", 'Bachelors'),
            (r'(?:diploma|associate)\s*(?:in|of)?\s*([A-Za-z\s]+)?', 'Diploma'),
        ]
        for pattern, deg_type in degree_patterns:
            for m in re.finditer(pattern, tl):
                field = m.group(1).strip().title() if m.group(1) else ""
                field = field[:60]
                ctx_start = max(0, m.start() - 150)
                ctx_end = min(len(tl), m.end() + 150)
                ctx = tl[ctx_start:ctx_end]
                institution = ""
                inst_patterns = [
                    r'(?:university|college|institute|school|academy)\s+(?:of\s+)?([A-Za-z\s]+)',
                    r'([A-Z][A-Za-z\s]+(?:university|college|institute|school|academy))',
                ]
                for ip in inst_patterns:
                    im = re.search(ip, text[ctx_start:ctx_end], re.IGNORECASE)
                    if im:
                        institution = im.group(0).strip().title()[:80]
                        break
                year = ""
                ym = re.search(r'(20[0-2]\d|19\d\d)', ctx)
                if ym: year = ym.group(1)
                gpa = ""
                gm = re.search(r'(?:gpa|cgpa|grade)[:\s]*(\d+\.?\d*)\s*(?:/\s*\d+\.?\d*)?', ctx)
                if gm: gpa = gm.group(0).strip()
                already = any(e['degreeType'] == deg_type and e.get('institution','') == institution for e in education)
                if not already:
                    education.append({
                        "degreeType": deg_type,
                        "field": field if field and len(field) > 2 else "Not specified",
                        "institution": institution or "Not specified",
                        "year": year,
                        "gpa": gpa,
                    })

        if not education and any(w in tl for w in ['degree','university','college','institute']):
            education.append({
                "degreeType": "Degree (type unclear)",
                "field": "Not specified",
                "institution": "Detected but couldn't parse",
                "year": "", "gpa": ""
            })
        return education

    def detect_certifications_detailed(self, text):
        """Find certifications with context"""
        tl = text.lower()
        found = []
        for cert in KNOWN_CERTIFICATIONS:
            if cert in tl:
                already = any(cert in fc['name'] or fc['name'] in cert for fc in found)
                if not already:
                    idx = tl.find(cert)
                    ctx_s = max(0, idx - 40)
                    ctx_e = min(len(tl), idx + len(cert) + 40)
                    snippet = text[ctx_s:ctx_e].strip().replace('\n', ' ')
                    found.append({"name": cert.title(), "context": f"...{snippet}..."})
        return found

    def detect_achievements_detailed(self, text):
        """Detect quantified achievements with type and context"""
        achievements = []
        patterns = [
            (r'(?:increased|improved|boosted|grew|raised|enhanced)\s+.*?\d+\s*%', 'percentage_increase'),
            (r'(?:reduced|decreased|lowered|cut|minimized|saved)\s+.*?\d+\s*%', 'percentage_decrease'),
            (r'\$\s*\d+[\d,]*(?:\.\d+)?\s*(?:million|billion|m|b|k)?', 'revenue_money'),
            (r'\d+[\d,]*\s*(?:users|customers|clients|visitors|downloads)', 'user_metrics'),
            (r'(?:managed|led|mentored|supervised)\s+(?:a\s+)?(?:team\s+of\s+)?\d+', 'team_size'),
            (r'\d+[%]\s*(?:increase|decrease|improvement|reduction|growth)', 'metric_change'),
            (r'\d+x\s+(?:faster|improvement|growth|increase)', 'multiplier'),
            (r'(?:99|100|98|97|96|95)(?:\.\d+)?[%]\s*(?:uptime|availability|accuracy)', 'reliability'),
            (r'\d+\+?\s+(?:projects|applications|systems|services|apis|features)\s+(?:delivered|completed|shipped|launched)', 'delivery'),
            (r'(?:awarded|award|recognition|prize|winner|patent)', 'award'),
        ]
        tl = text.lower()
        for p, category in patterns:
            for m in re.finditer(p, tl):
                snippet = m.group(0).strip()
                if snippet and snippet not in [a['text'] for a in achievements]:
                    idx = m.start()
                    line_start = text.rfind('\n', 0, idx) + 1
                    line_end = text.find('\n', idx)
                    if line_end == -1: line_end = min(idx + 150, len(text))
                    full_line = text[line_start:line_end].strip()
                    achievements.append({
                        "text": snippet,
                        "category": category.replace('_', ' ').title(),
                        "fullContext": full_line[:200],
                    })
        return achievements

    def get_quality_details(self, text, skills):
        """Detailed resume quality breakdown"""
        word_count = len(text.split())
        char_count = len(text)
        line_count = len(text.split('\n'))
        sections_detected = []
        section_keywords = {
            'Summary/Objective': ['summary','objective','profile','about me','professional summary'],
            'Experience': ['experience','employment','work history','professional experience','career'],
            'Education': ['education','academic','qualification','degree'],
            'Skills': ['skills','technical skills','technologies','competencies','expertise'],
            'Projects': ['projects','portfolio','personal projects','key projects'],
            'Certifications': ['certifications','certificates','credentials','licensed'],
            'Awards': ['awards','honors','achievements','recognition'],
            'Languages': ['languages','language proficiency'],
            'Contact': ['email','phone','linkedin','github','portfolio','address'],
        }
        tl = text.lower()
        for section, keywords in section_keywords.items():
            if any(kw in tl for kw in keywords):
                sections_detected.append(section)

        has_bullet_points = bool(re.search(r'[•\-\*]', text))
        has_numbers = bool(re.search(r'\d+[%$]|\$\d+|\d+\s*(?:users|projects|years)', tl))

        return {
            "wordCount": word_count,
            "characterCount": char_count,
            "lineCount": line_count,
            "skillsDetected": len(skills),
            "sectionsFound": sections_detected,
            "sectionCount": len(sections_detected),
            "hasBulletPoints": has_bullet_points,
            "hasQuantifiedResults": has_numbers,
            "lengthAssessment": "Good length" if 400 <= word_count <= 1200 else "Too short - add more detail" if word_count < 400 else "Consider trimming",
            "tips": [
                "Add more skills" if len(skills) < 8 else "Good skill coverage",
                "Add bullet points for readability" if not has_bullet_points else "Good use of bullet points",
                "Add numbers/metrics to achievements" if not has_numbers else "Good use of quantified results",
                f"Missing sections: {', '.join(set(section_keywords.keys()) - set(sections_detected))}" if len(sections_detected) < 5 else "All key sections present",
            ]
        }

    def get_seniority_details(self, resume_text, job_desc):
        """Detailed seniority comparison"""
        def detect_level(text):
            tl = text.lower()
            if any(w in tl for w in ['principal engineer','staff engineer','distinguished','vp of engineering','director of engineering','cto','chief technology','head of engineering']): return 'principal'
            elif any(w in tl for w in ['senior engineer','senior developer','senior software','lead engineer','lead developer','tech lead','team lead','senior analyst','sr.','sr ']): return 'senior'
            elif any(w in tl for w in ['mid-level','mid level','software engineer ii','developer ii','engineer ii']): return 'mid'
            elif any(w in tl for w in ['junior','associate','entry-level','entry level','software engineer i','developer i','graduate']): return 'junior'
            elif any(w in tl for w in ['intern','trainee','apprentice']): return 'intern'
            return 'unknown'

        c_level = detect_level(resume_text)
        j_level = detect_level(job_desc)
        level_names = {'intern':'Intern','junior':'Junior','mid':'Mid-Level','senior':'Senior','principal':'Principal/Staff','unknown':'Not Specified'}
        level_order = {'intern':0,'junior':1,'mid':2,'senior':3,'principal':4,'unknown':-1}
        c_ord = level_order.get(c_level, -1)
        j_ord = level_order.get(j_level, -1)

        if j_ord == -1 or c_ord == -1:
            match_desc = "Could not determine one or both levels"
            match_quality = "Neutral"
        elif c_ord == j_ord:
            match_desc = "Perfect seniority match"
            match_quality = "Excellent"
        elif c_ord == j_ord + 1:
            match_desc = "Candidate is one level above - overqualified but strong"
            match_quality = "Good"
        elif c_ord == j_ord - 1:
            match_desc = "Candidate is one level below - stretch role"
            match_quality = "Moderate"
        elif c_ord > j_ord:
            match_desc = "Candidate significantly overqualified"
            match_quality = "Concern"
        else:
            match_desc = "Candidate significantly underqualified for this level"
            match_quality = "Gap"

        return {
            "candidateLevel": level_names.get(c_level, 'Unknown'),
            "jobLevel": level_names.get(j_level, 'Unknown'),
            "candidateLevelRaw": c_level,
            "jobLevelRaw": j_level,
            "matchDescription": match_desc,
            "matchQuality": match_quality,
        }

    # ===================================================================
    # EXPERIENCE EXTRACTION (3 methods - unchanged)
    # ===================================================================
    def extract_experience(self, text):
        tl = text.lower()
        for p in [r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?experience',
                  r'experience[:\-\s]*(\d+)\+?\s*(?:years?|yrs?)',
                  r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:professional|work|industry|hands[\-\s]?on|relevant|total)\s*experience',
                  r'(?:over|more\s+than|around|nearly|about|~)\s*(\d+)\+?\s*(?:years?|yrs?)',
                  r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)\s*(?:in|with)',
                  r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:in\s+)?(?:software|it|tech|development|engineering|data|cloud|devops|web|mobile)',
                  r'total\s*(?:of\s+)?(\d+)\+?\s*(?:years?|yrs?)',
                  r'(?:bringing|with|having)\s+(\d+)\+?\s*(?:years?|yrs?)',
                  r'(?:experienced|seasoned)\s*(?:professional|developer|engineer)?\s*(?:with\s+)?(\d+)\+?\s*(?:years?|yrs?)']:
            m = re.search(p, tl)
            if m:
                try:
                    y = int(m.group(1))
                    if 0 < y < 50: return y
                except: pass
        yfd = self._calc_dates(text)
        if yfd > 0: return yfd
        yfp = self._est_positions(tl)
        if yfp > 0: return yfp
        return 0

    def _calc_dates(self, text):
        tl = text.lower(); cy, cm = datetime.now().year, datetime.now().month
        mp = r'(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
        pp = r'(?:present|current|now|ongoing|till\s*date|to\s*date|today)'
        ds = r'\s*(?:[-–—]+|\bto\b)\s*'
        ranges = []
        for m in re.finditer(rf'({mp})[\.,]?\s*((?:19|20)\d{{2}}){ds}(?:({mp})[\.,]?\s*((?:19|20)\d{{2}})|({pp}))', tl):
            try:
                sy=int(m.group(2));sm=self.month_map.get(m.group(1)[:3],1)
                if m.group(5):ey,em=cy,cm
                else:ey=int(m.group(4));em=self.month_map.get(m.group(3)[:3],12) if m.group(3) else 12
                if 1970<=sy<=cy and ey>=sy and not self._is_edu(tl,m.start(),m.end()):ranges.append((sy,sm,ey,em))
            except:pass
        for m in re.finditer(rf'(\d{{1,2}})\s*[/\-.]\s*((?:19|20)\d{{2}}){ds}(?:(\d{{1,2}})\s*[/\-.]\s*((?:19|20)\d{{2}})|({pp}))', tl):
            try:
                sm,sy=int(m.group(1)),int(m.group(2))
                if m.group(5):ey,em=cy,cm
                else:em,ey=int(m.group(3)),int(m.group(4))
                if 1<=sm<=12 and 1970<=sy<=cy and not self._is_edu(tl,m.start(),m.end()):ranges.append((sy,sm,ey,max(1,min(12,em))))
            except:pass
        if not ranges:
            for m in re.finditer(rf'\b((?:19|20)\d{{2}}){ds}(?:((?:19|20)\d{{2}})|({pp}))\b', tl):
                try:
                    sy=int(m.group(1));ey=cy if m.group(3) else int(m.group(2))
                    if 1970<=sy<=cy and ey>=sy and 0<ey-sy<=40 and not self._is_edu(tl,m.start(),m.end()):ranges.append((sy,1,ey,12))
                except:pass
        if ranges:
            tm=self._merge(ranges)
            if tm>=3:return min(max(round(tm/12),1),40)
        return 0

    def _is_edu(self, tl, s, e):
        edu=['university','college','school','institute','academy','bachelor','master','degree','diploma','graduated','gpa','cgpa','education','academic','b.tech','btech','b.e.','b.sc','bsc','m.tech','mtech','m.sc','msc','mba','ph.d','phd','doctorate','coursework','thesis']
        before=tl[max(0,s-300):s]
        for h in ['education','academic','qualification','schooling']:
            if h in before and not any(w in before[before.rfind(h):] for w in ['experience','employment','work history','professional','career']):return True
        return sum(1 for w in edu if w in tl[max(0,s-100):min(len(tl),e+80)])>=2

    def _merge(self, ranges):
        mr=sorted([(sy*12+sm,ey*12+em) for sy,sm,ey,em in ranges if ey*12+em>sy*12+sm])
        if not mr:return 0
        merged=[mr[0]]
        for s,e in mr[1:]:
            if s<=merged[-1][1]:merged[-1]=(merged[-1][0],max(merged[-1][1],e))
            else:merged.append((s,e))
        return sum(e-s for s,e in merged)

    def _est_positions(self, tl):
        pats=[r'(?:software|web|frontend|backend|full[\-\s]?stack|senior|junior|lead|staff|principal)\s+(?:engineer|developer|programmer)',r'(?:data|ml|ai|cloud|devops|platform|mobile|qa|test|automation)\s+(?:engineer|scientist|analyst|developer)',r'(?:project|product|engineering|program)\s+manager',r'(?:technical|solutions?|software)\s+architect',r'(?:business|systems?|data)\s+analyst',r'(?:scrum\s+master|tech(?:nical)?\s+lead)',r'\b(?:intern(?:ship)?|trainee)\b',r'\b(?:consultant|contractor|freelancer)\b']
        found=[]
        for p in pats:
            for m in re.findall(p,tl):
                if m not in found:found.append(m)
        n=len(found)
        if n>=5:return min(n*2,15)
        elif n>=3:return n*2
        elif n>=2:return 3
        elif n==1:return 1 if any(w in found[0] for w in ['intern','trainee']) else 2
        if any(w in tl for w in ['work experience','professional experience','employment history']):return 1
        return 0

    def extract_education(self, text):
        tl=text.lower()
        if any(w in tl for w in ['phd','ph.d','doctorate','doctoral']):return 15
        elif any(w in tl for w in ["master's",'masters','msc','m.sc','mba','ms degree','master of']):return 12
        elif any(w in tl for w in ["bachelor's",'bachelors','bsc','b.sc','b.tech','btech','b.e.','bachelor of']):return 10
        elif any(w in tl for w in ['diploma','associate','certification','bootcamp']):return 7
        elif any(w in tl for w in ['degree','university','college','institute']):return 8
        return 0

    def find_related_skills(self, missing, resume_skills):
        ml=missing.lower();rl=[s.lower() for s in resume_skills];c=0
        for gs in RELATED_SKILL_GROUPS.values():
            if ml in gs:c+=sum(1 for rs in rl if rs in gs and rs!=ml)
        return c

    # ===================================================================
    # 10-METRIC SCORING WITH FULL DETAILS
    # ===================================================================
    def calculate_score(self, matched, missing, resume_text, job_desc, exp_years, resume_skills):
        total = len(matched) + len(missing)

        # 1. Core Skill Match (25)
        skill_pct = len(matched) / total if total > 0 else 0
        skill_score = skill_pct * 25 if total > 0 else 12
        skill_details = self.get_skill_match_details(matched + missing, resume_text)

        # 2. Skill Depth (8)
        depth_details_list, avg_mentions = self.get_skill_depth_details(matched, resume_text)
        if not depth_details_list: depth_score = 0
        else:
            avg_depth = sum(float(d['depthScore'].split('/')[0]) for d in depth_details_list) / len(depth_details_list)
            depth_score = min(round((avg_depth / 3) * 8, 1), 8)

        # 3. Transferable (7)
        transferable_details = self.get_transferable_details(missing, resume_skills)
        related_bonus = 0
        if missing:
            tr = sum(min(self.find_related_skills(m, resume_skills), 3) for m in missing)
            mp = len(missing) * 2
            if mp > 0: related_bonus = min((tr / mp) * 7, 7)

        # 4. TF-IDF Cosine (10)
        tfidf_details = self.get_tfidf_details(resume_text, job_desc)
        cosine_score = min((tfidf_details['similarityPercent'] / 100) * 10, 10)

        # 5. Keywords (5)
        keyword_details = self.get_keyword_details(resume_text, job_desc)
        keyword_score = (keyword_details['matchPercent'] / 100) * 5

        # 6. Experience (15)
        req_exp = 0
        for p in [r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)', r'(?:minimum|at least)\s*(\d+)\s*(?:years?|yrs?)']:
            m = re.search(p, job_desc.lower())
            if m:
                try: req_exp = int(m.group(1))
                except: pass
                break
        if exp_years > 0:
            if req_exp > 0:
                er = min(exp_years / req_exp, 1.5)
                exp_score = 15 if er>=1.0 else 12 if er>=0.75 else 9 if er>=0.5 else 6 if er>=0.25 else 3
            else:
                exp_score = 15 if exp_years>=10 else 13 if exp_years>=7 else 11 if exp_years>=5 else 9 if exp_years>=3 else 7 if exp_years>=2 else 5 if exp_years>=1 else 3
        else: exp_score = 0
        experience_entries = self.extract_work_experiences(resume_text)

        # 7. Education (7)
        edu_raw = self.extract_education(resume_text)
        edu_score = min(round(edu_raw * 0.7), 7)
        education_entries = self.extract_education_details(resume_text)

        # 8. Certifications (5)
        cert_entries = self.detect_certifications_detailed(resume_text)
        cert_score = 5 if len(cert_entries)>=4 else 4 if len(cert_entries)>=3 else 3 if len(cert_entries)>=2 else 2 if len(cert_entries)>=1 else 0

        # 9. Achievements (8)
        ach_entries = self.detect_achievements_detailed(resume_text)
        ach_score = 8 if len(ach_entries)>=6 else 6 if len(ach_entries)>=4 else 5 if len(ach_entries)>=3 else 4 if len(ach_entries)>=2 else 2 if len(ach_entries)>=1 else 0

        # 10. Quality (5)
        quality_details = self.get_quality_details(resume_text, resume_skills)
        rl = len(resume_text)
        quality = 5 if rl>=3000 else 4 if rl>=2000 else 3 if rl>=1000 else 2 if rl>=500 else 1
        if len(resume_skills) >= 15: quality = min(quality+1, 5)
        elif len(resume_skills) >= 10: quality = min(quality+0.5, 5)

        # 11. Seniority (5)
        seniority_details = self.get_seniority_details(resume_text, job_desc)
        level_order = {'intern':0,'junior':1,'mid':2,'senior':3,'principal':4,'unknown':-1}
        c_lvl = level_order.get(seniority_details['candidateLevelRaw'], -1)
        j_lvl = level_order.get(seniority_details['jobLevelRaw'], -1)
        if j_lvl == -1 or c_lvl == -1: seniority_score = 3
        elif c_lvl == j_lvl: seniority_score = 5
        elif c_lvl == j_lvl + 1: seniority_score = 4
        elif c_lvl == j_lvl - 1: seniority_score = 3
        elif c_lvl > j_lvl: seniority_score = 3
        else: seniority_score = max(0, 5 - (j_lvl - c_lvl) * 2)

        overall = max(5, min(98, round(
            skill_score + depth_score + related_bonus + cosine_score + keyword_score +
            exp_score + edu_score + cert_score + ach_score + quality + seniority_score
        )))

        return overall, {
            "total": overall, "totalMax": 100,
            # Scores
            "coreSkillMatch": round(skill_score, 1), "coreSkillMatchMax": 25,
            "skillDepth": round(depth_score, 1), "skillDepthMax": 8,
            "transferableSkills": round(related_bonus, 1), "transferableSkillsMax": 7,
            "cosineSimilarity": round(cosine_score, 1), "cosineSimilarityMax": 10,
            "keywordOverlap": round(keyword_score, 1), "keywordOverlapMax": 5,
            "experience": exp_score, "experienceMax": 15,
            "education": edu_score, "educationMax": 7,
            "certifications": cert_score, "certificationsMax": 5,
            "achievements": ach_score, "achievementsMax": 8,
            "resumeQuality": round(quality, 1), "resumeQualityMax": 5,
            "seniorityMatch": seniority_score, "seniorityMatchMax": 5,
            # Percentages
            "skillMatchPct": round(skill_pct * 100, 1),
            "keywordOverlapPct": round(keyword_details['matchPercent'], 1),
            "cosineSimilarityPct": round(tfidf_details['similarityPercent'], 1),
            "candidateExperience": exp_years, "requiredExperience": req_exp,
            # ===== DETAILED BREAKDOWNS (for tap-to-view) =====
            "details": {
                "coreSkillMatch": {
                    "description": "Which required skills were found in your resume",
                    "matchedCount": len(matched), "missingCount": len(missing), "totalRequired": total,
                    "skills": skill_details,
                },
                "skillDepth": {
                    "description": "How deeply each matched skill is demonstrated",
                    "averageMentions": avg_mentions,
                    "skills": depth_details_list,
                },
                "transferableSkills": {
                    "description": "Related skills you have for each missing skill",
                    "skills": transferable_details,
                },
                "cosineSimilarity": {
                    "description": "AI-powered semantic similarity between resume and JD",
                    **tfidf_details,
                },
                "keywordOverlap": {
                    "description": "Direct keyword matches between resume and job description",
                    **keyword_details,
                },
                "experience": {
                    "description": "Work experience detected from your resume",
                    "totalYears": exp_years,
                    "requiredYears": req_exp,
                    "meetsRequirement": exp_years >= req_exp if req_exp > 0 else True,
                    "positions": experience_entries,
                    "detectionMethod": "explicit" if exp_years > 0 else "none",
                },
                "education": {
                    "description": "Educational qualifications found in your resume",
                    "entries": education_entries,
                    "highestDegree": education_entries[0]['degreeType'] if education_entries else "Not detected",
                },
                "certifications": {
                    "description": "Professional certifications detected",
                    "count": len(cert_entries),
                    "entries": cert_entries,
                    "suggestion": "Great certification portfolio!" if len(cert_entries) >= 3 else f"Consider getting certified to boost your profile" if len(cert_entries) == 0 else "Good start - consider more certifications",
                },
                "achievements": {
                    "description": "Quantified achievements and measurable impact",
                    "count": len(ach_entries),
                    "entries": ach_entries,
                    "suggestion": "Excellent use of metrics!" if len(ach_entries) >= 4 else "Add more quantified results (%, $, numbers)" if len(ach_entries) < 2 else "Good - try to add more specific numbers",
                },
                "resumeQuality": {
                    "description": "Overall resume structure and quality",
                    **quality_details,
                },
                "seniorityMatch": {
                    "description": "How your seniority level matches the job requirements",
                    **seniority_details,
                },
            }
        }

    def generate_analysis(self, matched, missing, score, exp, company, job_data, sb):
        total = len(matched) + len(missing)
        pct = (len(matched) / total * 100) if total > 0 else 0
        level = "Senior" if exp >= 7 else "Mid-Level" if exp >= 4 else "Junior" if exp >= 1 else "Entry-Level"

        if score >= 80: rec, rr = "STRONG HIRE", f"Excellent ({score}/100)"
        elif score >= 65: rec, rr = "HIRE", f"Good ({score}/100)"
        elif score >= 50: rec, rr = "MAYBE", f"Moderate ({score}/100)"
        elif score >= 35: rec, rr = "CONDITIONAL", f"Below average ({score}/100)"
        else: rec, rr = "NEEDS DEVELOPMENT", f"Low ({score}/100)"

        summary = f"{score}/100 - {'Excellent match!' if score>=80 else 'Good candidate' if score>=65 else 'Decent fit' if score>=50 else 'Partial match' if score>=35 else 'Needs development'}"
        fit = f"{len(matched)}/{total} skills ({int(pct)}% match). "
        if matched: fit += f"Strong: {', '.join(matched[:5])}. "
        if missing: fit += f"Needs: {', '.join(missing[:4])}. "
        fit += f"{exp}y exp ({level}). Score: {score}/100."

        strengths = []
        if matched: strengths.append(f"Proficient in {len(matched)} required skills: {', '.join(matched[:6])}")
        if sb.get('skillDepth', 0) >= 5: strengths.append("Deep expertise demonstrated through projects")
        if exp >= 5: strengths.append(f"Strong: {exp} years experience")
        elif exp >= 2: strengths.append(f"{exp} years experience")
        elif exp >= 1: strengths.append(f"{exp} year experience")
        certs = sb.get('details', {}).get('certifications', {}).get('count', 0)
        if certs >= 2: strengths.append(f"{certs} relevant certifications")
        achs = sb.get('details', {}).get('achievements', {}).get('count', 0)
        if achs >= 3: strengths.append(f"{achs} quantified achievements showing impact")
        if sb.get('cosineSimilarityPct', 0) >= 30: strengths.append("Resume aligns well with JD")
        if sb.get('education', 0) >= 6: strengths.append("Strong educational background")
        if pct >= 70: strengths.append("Matches majority of required skills")
        if not strengths: strengths.append("Shows willingness to learn")

        weaknesses = []
        if missing: weaknesses.append(f"Missing {len(missing)} skills: {', '.join(missing[:4])}")
        if exp == 0: weaknesses.append("No professional experience detected")
        if achs == 0: weaknesses.append("No quantified achievements - add metrics")
        if certs == 0: weaknesses.append("No certifications - consider getting certified")
        if sb.get('cosineSimilarityPct', 0) < 15: weaknesses.append("Resume doesn't align well with JD - tailor it")
        if sb.get('skillDepth', 0) < 3: weaknesses.append("Skills mentioned but not demonstrated in projects")
        if not weaknesses: weaknesses.append("Minor areas for growth")

        tips = []
        if matched: tips.append(f"Highlight deep experience with {', '.join(matched[:3])}")
        tips.append("Add quantified achievements (e.g., 'Improved performance by 40%')")
        tips.append(f"Research {company or 'the company'}'s tech stack")
        if missing: tips.append(f"Show you're learning {missing[0]}")
        tips.append("Practice STAR method for behavioral questions")

        if missing:
            tw = min(max(len(missing) * 3, 4), 24)
            lp = {"totalTimeEstimate": f"{tw-2}-{tw+2} weeks", "weeklyCommitment": "10-15h",
                  "priorityOrder": missing[:6], "careerImpact": f"Can increase score to 85-95",
                  "salaryPotential": f"+${len(missing)*5}K-${len(missing)*10}K",
                  "jobReadiness": f"Ready in {max(tw//2,2)} weeks",
                  "milestones": [{"week":2,"goal":f"Complete {missing[0]} basics","status":"upcoming"},
                                 {"week":4,"goal":"Portfolio project","status":"upcoming"},
                                 {"week":tw,"goal":"Fully qualified","status":"upcoming"}]}
        else:
            lp = {"totalTimeEstimate":"Ready now!","weeklyCommitment":"Interview prep","priorityOrder":[],
                  "careerImpact":"Already qualified","salaryPotential":"Strong position",
                  "jobReadiness":"Immediately ready","milestones":[]}

        gaps = []
        for i, sk in enumerate(missing[:6]):
            sd = get_skill_resources_data(sk); imp = max(95-i*7, 55)
            gaps.append({
                "skill":sk,"displayName":safe_str(sd.get('displayName'),sk.title()),
                "category":safe_str(sd.get('category'),'Technical'),
                "importancePercentage":imp,
                "priorityLevel":"Critical" if i<2 else "High" if i<4 else "Medium",
                "learningUrgency":"Start immediately" if i<2 else "Start this week" if i<4 else "Within 2 weeks",
                "difficulty":safe_str(sd.get('difficulty'),'Medium'),
                "learningTime":safe_str(sd.get('learningTime'),'4-6 weeks'),
                "demandLevel":safe_str(sd.get('demandLevel'),'High'),
                "salaryImpact":safe_str(sd.get('salaryImpact'),'+$10K'),
                "description":safe_str(sd.get('description')),
                "whyLearn":f"{sk.title()} is required for this role.",
                "careerPaths":safe_list(sd.get('careerPaths')),
                "relatedSkills":safe_list(sd.get('relatedSkills'))[:4],
                "topCourses":safe_list(sd.get('freeCoursePlatforms'))[:3],
                "topVideos":safe_list(sd.get('youtubeChannels'))[:3],
                "topPractice":safe_list(sd.get('practiceWebsites'))[:2],
                "topProjects":safe_list(sd.get('projectIdeas'))[:3],
                "roadmap":safe_dict(sd.get('learningRoadmap')),
                "certifications":safe_list(sd.get('certifications'))
            })

        return {"overallScore":score,"experienceLevel":level,"fitAnalysis":fit,
                "jobMatchSummary":summary,"hiringRecommendation":rec,"recommendationReason":rr,
                "strengths":strengths,"weaknesses":weaknesses,"interviewTips":tips,
                "overallLearningPlan":lp,"skillGapAnalysis":gaps}


analyzer = ResumeAnalyzer(client if AI_AVAILABLE else None)


# ============================================================================
# ROUTES
# ============================================================================
@app.route('/api/auth/signup', methods=['POST'])
def signup():
    try:
        d=request.json;email=d.get('email','').lower().strip();pw=d.get('password','');name=d.get('name','').strip()
        if not all([email,pw,name]):return jsonify({"error":"All fields required"}),400
        if len(pw)<6:return jsonify({"error":"Password min 6 chars"}),400
        if DatabaseManager.get_user_by_email(email):return jsonify({"error":"Email already registered"}),400
        u=DatabaseManager.create_user(email,name,pw);t=DatabaseManager.create_session(u.id)
        return jsonify({"message":"Signup successful","user":{"email":u.email,"name":u.name},"token":t}),201
    except Exception as e:return jsonify({"error":str(e)}),500

@app.route('/api/auth/login', methods=['POST'])
def login():
    try:
        d=request.json;email=d.get('email','').lower().strip();pw=d.get('password','')
        if not all([email,pw]):return jsonify({"error":"Email and password required"}),400
        u=DatabaseManager.get_user_by_email(email)
        if not u or not check_password_hash(u.password,pw):return jsonify({"error":"Invalid credentials"}),401
        t=DatabaseManager.create_session(u.id)
        return jsonify({"message":"Login successful","user":{"email":u.email,"name":u.name},"token":t}),200
    except Exception as e:return jsonify({"error":str(e)}),500

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    DatabaseManager.delete_session(request.headers.get('Authorization','').replace('Bearer ',''))
    return jsonify({"message":"Logged out"}),200

@app.route('/api/auth/me', methods=['GET'])
def get_current_user():
    u=DatabaseManager.get_user_by_token(request.headers.get('Authorization','').replace('Bearer ',''))
    if not u:return jsonify({"error":"Not authenticated"}),401
    return jsonify({"user":{"email":u.email,"name":u.name}}),200

@app.route('/api/analyze', methods=['POST'])
def analyze_resume():
    try:
        u=DatabaseManager.get_user_by_token(request.headers.get('Authorization','').replace('Bearer ',''))
        if not u:return jsonify({"error":"Please login"}),401
        if 'resume' not in request.files:return jsonify({"error":"No resume uploaded"}),400
        file=request.files['resume']
        jd=request.form.get('jobDescription','').strip()
        company=request.form.get('companyName','').strip()
        if not jd or len(jd)<50:return jsonify({"error":"Job description too short (min 50 chars)"}),400

        if file.filename.lower().endswith('.pdf'):rt=analyzer.extract_pdf_text(file)
        elif file.filename.lower().endswith('.docx'):rt=analyzer.extract_docx_text(file)
        else:return jsonify({"error":"Only PDF/DOCX supported"}),400
        if len(rt)<100:return jsonify({"error":"Could not extract resume text."}),400

        js=analyzer.extract_skills_from_text(jd)
        if not js:return jsonify({"error":"No technical skills found in job description."}),400

        matched,missing=analyzer.match_skills(js,rt)
        rs=analyzer.extract_skills_from_text(rt)
        exp=analyzer.extract_experience(rt)
        score,sb=analyzer.calculate_score(matched,missing,rt,jd,exp,rs)
        analysis=analyzer.generate_analysis(matched,missing,score,exp,company,{"required_skills":js},sb)

        lr=[]
        for sk in missing[:6]:
            if is_soft_skill(sk):continue
            sd=get_skill_resources_data(sk)
            lr.append({
                "skillName":sk,"displayName":safe_str(sd.get('displayName'),sk.title()),
                "category":safe_str(sd.get('category'),'Technical'),
                "difficulty":safe_str(sd.get('difficulty'),'Medium'),
                "learningTime":safe_str(sd.get('learningTime'),'4-6 weeks'),
                "salaryImpact":safe_str(sd.get('salaryImpact'),'+$10K'),
                "demandLevel":safe_str(sd.get('demandLevel'),'High'),
                "description":safe_str(sd.get('description')),
                "careerPaths":safe_list(sd.get('careerPaths')),
                "relatedSkills":safe_list(sd.get('relatedSkills')),
                "courses":safe_list(sd.get('freeCoursePlatforms')),
                "youtubeVideos":safe_list(sd.get('youtubeChannels')),
                "practiceResources":safe_list(sd.get('practiceWebsites')),
                "projects":safe_list(sd.get('projectIdeas')),
                "roadmap":safe_dict(sd.get('learningRoadmap')),
                "certifications":safe_list(sd.get('certifications'))
            })

        result={
            "message":"Analysis complete","companyName":safe_str(company,"Not specified"),
            "overallScore":safe_int(analysis['overallScore']),
            "jobMatchScore":safe_int(sb.get('skillMatchPct',0)),
            "textSimilarity":safe_int(sb.get('cosineSimilarityPct',0)),
            "experienceLevel":safe_str(analysis['experienceLevel']),
            "fitAnalysis":safe_str(analysis['fitAnalysis']),
            "jobMatchSummary":safe_str(analysis['jobMatchSummary']),
            "hiringRecommendation":safe_str(analysis['hiringRecommendation']),
            "recommendationReason":safe_str(analysis['recommendationReason']),
            "strengths":safe_list(analysis['strengths']),
            "weaknesses":safe_list(analysis['weaknesses']),
            "interviewTips":safe_list(analysis['interviewTips']),
            "jobRequirements":safe_list(js),"requiredSkills":safe_list(js),"preferredSkills":[],
            "candidateSkills":safe_list(rs),"matchedSkills":safe_list(matched),
            "missingSkills":safe_list(missing),
            "candidateExperience":safe_int(exp),
            "requiredExperience":safe_int(sb.get('requiredExperience',0)),
            "scoreBreakdown":sb,
            "overallLearningPlan":safe_dict(analysis['overallLearningPlan']),
            "skillGapAnalysis":safe_list(analysis['skillGapAnalysis']),
            "learningResources":lr,
            "skillsFound":safe_list(rs),"skillsRequired":safe_list(js),
            "skillGaps":safe_list(missing),
            "similarityScore":safe_int(sb.get('cosineSimilarityPct',0)),
        }

        DatabaseManager.create_analysis(u.id,company,result)
        return jsonify(result),200

    except Exception as e:
        import traceback;traceback.print_exc()
        return jsonify({"error":str(e)}),500

@app.route('/api/health', methods=['GET'])
def health():
    try:
        return jsonify({"status":"healthy","database":"MySQL","users":User.query.count(),
            "skills":len([s for s in KNOWN_SKILLS if not is_soft_skill(s)]),
            "scoring":"10-metric with tap-to-view details",
            "metrics":["coreSkillMatch","skillDepth","transferableSkills","cosineSimilarity",
                       "keywordOverlap","experience","education","certifications","achievements",
                       "resumeQuality","seniorityMatch"]}),200
    except Exception as e:return jsonify({"status":"error","message":str(e)}),500

if __name__=='__main__':
    print("="*60)
    print("RESUME ANALYZER - 10 METRICS + TAP-TO-VIEW DETAILS")
    print("="*60)
    print("Each metric returns detailed breakdown when tapped:")
    print("  1. Core Skills    → which matched/missing + context snippets")
    print("  2. Skill Depth    → mention count + project context per skill")
    print("  3. Transferable   → related skills you have for each gap")
    print("  4. TF-IDF Cosine  → top shared terms + interpretation")
    print("  5. Keywords       → matched/unmatched keyword lists")
    print("  6. Experience     → each job title, company, dates, duration")
    print("  7. Education      → degree, institution, field, year, GPA")
    print("  8. Certifications → each cert found with context")
    print("  9. Achievements   → each quantified result with category")
    print("  10. Quality       → word count, sections, bullet points")
    print("  11. Seniority     → candidate vs job level comparison")
    print("="*60)
    with app.app_context():
        try:db.create_all();print("Database ready")
        except Exception as e:print(f"DB Warning: {e}")
    print(f"Server: http://localhost:5000")
    print("="*60)
    app.run(debug=True,port=5000)